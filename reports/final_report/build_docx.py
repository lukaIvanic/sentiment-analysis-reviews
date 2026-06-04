"""Build and polish the final DOCX report.

The Markdown source is the canonical text. Pandoc handles Markdown parsing, and
this script applies the Word-level formatting that Pandoc does not reliably set:
figure centering, table borders/geometry, paragraph rhythm, and report styles.
"""

from __future__ import annotations

import argparse
import re
import shutil
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
REPORT_DIR = ROOT / "reports" / "final_report"
DEFAULT_SOURCE = REPORT_DIR / "report.md"
DEFAULT_OUTPUT = REPORT_DIR / "analiza_sentimenta_recenzija.docx"
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
APPENDIX_H_CARD_HEADING_RE = re.compile(r"^H\.(1|3|4|5|6|7|8|9|10|11)\b")


def require_pandoc() -> str:
    pandoc = shutil.which("pandoc")
    if not pandoc:
        raise RuntimeError("pandoc is required to build the DOCX report")
    return pandoc


def run_pandoc(source: Path, temp_docx: Path) -> None:
    pandoc = require_pandoc()
    cmd = [
        pandoc,
        str(source),
        "--from",
        "markdown+pipe_tables+link_attributes",
        "--to",
        "docx",
        "--resource-path",
        f"{ROOT}:{REPORT_DIR}",
        "--metadata",
        "link-citations=false",
        "-o",
        str(temp_docx),
    ]
    subprocess.run(cmd, cwd=ROOT, check=True)


def set_font(run, *, name: str = "Calibri", size: float | None = None,
             color: str | None = None, bold: bool | None = None,
             italic: bool | None = None) -> None:
    font = run.font
    font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        font.size = Pt(size)
    if color is not None:
        font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        font.bold = bold
    if italic is not None:
        font.italic = italic


def configure_styles(doc: Document) -> None:
    style_tokens = {
        "Normal": ("Calibri", 11, "111827", 0, 6, 1.10),
        "Body Text": ("Calibri", 11, "111827", 0, 6, 1.10),
        "First Paragraph": ("Calibri", 11, "111827", 0, 6, 1.10),
        "Compact": ("Calibri", 10.5, "111827", 0, 4, 1.10),
        "Image Caption": ("Calibri", 9, "475569", 0, 10, 1.05),
        "Caption": ("Calibri", 9, "475569", 0, 10, 1.05),
        "Source Code": ("Consolas", 9, "111827", 0, 4, 1.05),
        "Block Text": ("Calibri", 10.5, "1F2937", 3, 8, 1.10),
    }
    for name, (font_name, size, color, before, after, line_spacing) in style_tokens.items():
        if name not in doc.styles:
            continue
        style = doc.styles[name]
        style.font.name = font_name
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        pf = style.paragraph_format
        pf.space_before = Pt(before)
        pf.space_after = Pt(after)
        pf.line_spacing = line_spacing

    heading_tokens = {
        "Heading 1": (16, "2E74B5", 16, 8),
        "Heading 2": (13, "2E74B5", 12, 6),
        "Heading 3": (12, "1F4D78", 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        if name not in doc.styles:
            continue
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        pf = style.paragraph_format
        pf.space_before = Pt(before)
        pf.space_after = Pt(after)
        pf.line_spacing = 1.10
        pf.keep_with_next = True

    title_styles = {
        "Title": (22, "0B2545", True, 0, 4),
        "Subtitle": (12, "475569", False, 0, 4),
        "Author": (10.5, "475569", False, 0, 2),
        "Date": (10.5, "475569", False, 0, 12),
    }
    for name, (size, color, bold, before, after) in title_styles.items():
        if name not in doc.styles:
            continue
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor.from_string(color)
        pf = style.paragraph_format
        pf.space_before = Pt(before)
        pf.space_after = Pt(after)
        pf.line_spacing = 1.05


def configure_sections(doc: Document) -> None:
    for section in doc.sections:
        section.start_type = WD_SECTION.NEW_PAGE
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)


def paragraph_has_drawing(paragraph) -> bool:
    return bool(paragraph._element.xpath(".//w:drawing"))


def should_start_new_page(paragraph) -> bool:
    text = paragraph.text.strip()
    return bool(APPENDIX_H_CARD_HEADING_RE.match(text)) or text.startswith("Dodatak I:")


def polish_paragraphs(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style is not None else ""
        pf = paragraph.paragraph_format

        if style_name in {"Title", "Subtitle", "Author", "Date"}:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        if paragraph_has_drawing(paragraph):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.space_before = Pt(8)
            pf.space_after = Pt(2)
            pf.keep_with_next = True
            continue

        if style_name in {"Image Caption", "Caption"}:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.space_before = Pt(0)
            pf.space_after = Pt(10)
            pf.line_spacing = 1.05
            pf.keep_together = True
            for run in paragraph.runs:
                set_font(run, size=9, color="475569", italic=True)
            continue

        if style_name.startswith("Heading"):
            pf.keep_with_next = True
            if should_start_new_page(paragraph):
                pf.page_break_before = True
            continue

        if style_name in {"Body Text", "First Paragraph", "Normal"}:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.space_after = Pt(6)
            pf.line_spacing = 1.10

        for run in paragraph.runs:
            if style_name == "Source Code":
                set_font(run, name="Consolas", size=9, color="111827")
            elif run.text.strip():
                set_font(run, size=11)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120) -> None:
    tbl_pr = table._tbl.tblPr
    existing = tbl_pr.find(qn("w:tblCellMar"))
    if existing is not None:
        tbl_pr.remove(existing)
    margins = OxmlElement("w:tblCellMar")
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        margins.append(node)
    tbl_pr.append(margins)


def set_table_borders(table, color="CBD5E1", size="6") -> None:
    tbl_pr = table._tbl.tblPr
    existing = tbl_pr.find(qn("w:tblBorders"))
    if existing is not None:
        tbl_pr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)
        borders.append(node)
    tbl_pr.append(borders)


def set_table_width(table, widths: list[int]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for col in list(grid):
        grid.remove(col)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def choose_widths(table) -> list[int]:
    cols = len(table.columns)
    first_row = [cell.text.strip().lower() for cell in table.rows[0].cells] if table.rows else []

    if cols == 7 and "model" in first_row[0]:
        return [2200, 1120, 1120, 1320, 1180, 1180, 1240]
    if cols == 5 and "model" in first_row[0]:
        return [3000, 1590, 1590, 1590, 1590]
    if cols == 4 and "model" in first_row[0]:
        return [3000, 2120, 2120, 2120]
    if cols == 4:
        return [2300, 3600, 1700, 1760]
    if cols == 3 and "komanda" in " ".join(first_row):
        return [2100, 5300, 1960]
    if cols == 3:
        return [3000, 1700, 4660]
    if cols == 2:
        return [2700, 6660]
    return [CONTENT_WIDTH_DXA // cols] * cols


def polish_tables(doc: Document) -> None:
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        set_table_width(table, choose_widths(table))
        set_table_borders(table)
        set_cell_margins(table)

        for row_idx, row in enumerate(table.rows):
            tr_pr = row._tr.get_or_add_trPr()
            if tr_pr.find(qn("w:cantSplit")) is None:
                tr_pr.append(OxmlElement("w:cantSplit"))
            if row_idx == 0 and tr_pr.find(qn("w:tblHeader")) is None:
                tbl_header = OxmlElement("w:tblHeader")
                tbl_header.set(qn("w:val"), "true")
                tr_pr.append(tbl_header)
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                if row_idx == 0:
                    set_cell_shading(cell, "F2F4F7")
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    paragraph.paragraph_format.space_after = Pt(2)
                    paragraph.paragraph_format.line_spacing = 1.05
                    for run in paragraph.runs:
                        set_font(
                            run,
                            size=8.6 if len(table.columns) >= 6 else 9.2,
                            color="111827",
                            bold=True if row_idx == 0 else None,
                        )


def add_table_spacers(doc: Document) -> None:
    for table in doc.tables:
        spacer = OxmlElement("w:p")
        p_pr = OxmlElement("w:pPr")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), "80")
        spacing.set(qn("w:after"), "80")
        p_pr.append(spacing)
        spacer.append(p_pr)

        run = OxmlElement("w:r")
        r_pr = OxmlElement("w:rPr")
        size = OxmlElement("w:sz")
        size.set(qn("w:val"), "2")
        r_pr.append(size)
        text = OxmlElement("w:t")
        text.set(qn("xml:space"), "preserve")
        text.text = " "
        run.append(r_pr)
        run.append(text)
        spacer.append(run)
        table._tbl.addnext(spacer)


def resize_images(doc: Document) -> None:
    max_width = Inches(6.15)
    for shape in doc.inline_shapes:
        if shape.width and shape.height:
            if shape.width > max_width:
                ratio = max_width / shape.width
                shape.width = max_width
                shape.height = int(shape.height * ratio)


def add_footer(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer
        if footer.paragraphs:
            paragraph = footer.paragraphs[0]
        else:
            paragraph = footer.add_paragraph()
        paragraph.text = "Analiza sentimenta recenzija filmova"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(4)
        for run in paragraph.runs:
            set_font(run, size=8.5, color="64748B")


def polish_docx(temp_docx: Path, output: Path) -> None:
    doc = Document(temp_docx)
    configure_sections(doc)
    configure_styles(doc)
    polish_paragraphs(doc)
    polish_tables(doc)
    add_table_spacers(doc)
    resize_images(doc)
    add_footer(doc)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def build(source: Path, output: Path) -> None:
    with tempfile.TemporaryDirectory(prefix="sentiment-report-") as tmp:
        temp_docx = Path(tmp) / "pandoc.docx"
        run_pandoc(source, temp_docx)
        polish_docx(temp_docx, output)


def main() -> None:
    parser = argparse.ArgumentParser(description="Build the polished final DOCX report.")
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()
    build(args.source, args.output)
    print(f"Built {args.output}")


if __name__ == "__main__":
    main()
